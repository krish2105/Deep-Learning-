#!/usr/bin/env python3
"""Generate the group contribution statement.

A separate signed document, because "Team Collaboration" and "equal
contribution" are graded criteria and a title-page table is too thin to
evidence them. Each member's section names what they owned, the specific
artefacts they produced, and — more usefully — a decision they made that
changed the project.

Run:  python scripts/build_contributions.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).resolve().parent.parent / "deliverables" / "report"
OUT.mkdir(parents=True, exist_ok=True)

INSTRUMENT = RGBColor(0x1B, 0x7D, 0x97)
INK = RGBColor(0x11, 0x16, 0x1A)
MUTED = RGBColor(0x5C, 0x66, 0x6D)

doc = Document()
for s in doc.sections:
    s.top_margin = s.bottom_margin = Cm(2.2)
    s.left_margin = s.right_margin = Cm(2.4)

n = doc.styles["Normal"]
n.font.name = "Calibri"
n.font.size = Pt(10.5)
n.font.color.rgb = INK
n.paragraph_format.space_after = Pt(6)
n.paragraph_format.line_spacing = 1.2


def para(t="", *, size=10.5, bold=False, italic=False, colour=None, align=None, after=None):
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    r.font.color.rgb = colour or INK
    if align is not None:
        p.alignment = align
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    return p


def shade(cell, hexc):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hexc)
    cell._tc.get_or_add_tcPr().append(el)


def table(headers, rows, widths=None, font=9):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(font)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade(c, "1B7D97")
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            rr = cells[i].paragraphs[0].add_run(str(v))
            rr.font.size = Pt(font)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    return t


def bullets(items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(it).font.size = Pt(10)


# ═══════════════════════════════════════════════════════════════════════
para("Statement of Contribution", size=20, bold=True, colour=INSTRUMENT, after=2)
para("SENTINEL-CXR — Uncertainty-Aware Chest Radiograph Triage", size=12, after=2)
para("Final Group Project · Deep Learning (MAIB AI 114)", size=10, colour=MUTED, after=1)
para("Prof Anshul Gupta · S P Jain School of Global Management, Dubai",
     size=10, colour=MUTED, after=14)

para(
    "The three members contributed equally in effort and each owned a distinct "
    "vertical of the system end to end. Work was divided by layer rather than "
    "by task so that every member carried a component from design through "
    "implementation to deployment, and could defend it independently. Design "
    "decisions, the evaluation protocol and the ethical analysis were agreed "
    "jointly before implementation began.",
    after=12,
)

table(
    ["Member", "Student ID", "Ownership"],
    [
        ["Krishna Mathur", "AS25DXB018", "Models, calibration and explainability"],
        ["Atharva Soundankar", "AS25DXB020", "Backend, deployment and evaluation"],
        ["Yash Petkar", "AS25DXB021", "Interface, dashboard and fairness reporting"],
    ],
    widths=[4.5, 3.0, 8.5],
)
doc.add_paragraph()

MEMBERS = [
    (
        "Krishna Mathur — AS25DXB018",
        "Models, calibration and explainability",
        [
            "Classification backbone: DenseNet-121 adaptation to single-channel input, "
            "preserving the pretrained RGB filters by summing rather than discarding them.",
            "ONNX export and int8 quantisation of the deployed classifier, including the "
            "verification harness that refuses to ship a model which cannot execute.",
            "Split conformal implementation with the finite-sample correction, the "
            "abstention rules, and the calibration run over 4,999 real radiographs.",
            "Class activation mapping computed inside the exported graph, giving the "
            "orchestrator explainability without a backward pass.",
            "Uncertainty decomposition into aleatoric and epistemic components, and the "
            "test-time-augmentation fallback.",
            "The eleven training notebooks covering weeks 1–12 of the unit.",
        ],
        "Judgement call: on discovering that the published checkpoint contains no dropout "
        "layers, rejected the option of injecting dropout at inference — which would have "
        "produced a number that looked like a Bayesian posterior but was not one — in "
        "favour of test-time augmentation, with the method used reported in every response.",
    ),
    (
        "Atharva Soundankar — AS25DXB020",
        "Backend, deployment and evaluation",
        [
            "FastAPI orchestration service: authentication, study lifecycle, audit trail, "
            "and the analysis pipeline that sequences gate, classifier, conformal head, "
            "triage and reporting.",
            "Three-tier deployment across Vercel, Render and Hugging Face Spaces, and the "
            "graceful-degradation path that keeps the system diagnosing when the optional "
            "inference core is unavailable.",
            "The evaluation harness and CI safety gates: coverage must not regress and no "
            "pathology may escape the report-grounding filter.",
            "Resilience work: per-client rate limiting behind a proxy, self-reissuing demo "
            "sessions, and the diagnostics that report why a subsystem failed.",
            "83 automated tests across the conformal head, uncertainty, triage ordering, "
            "progression logic, grounding and the full API surface.",
        ],
        "Judgement call: rather than treat a cold inference service as an outage, "
        "designed the system to answer immediately on a local quantised model and label "
        "the response as reduced — on the principle that a clinical user must always know "
        "which mode produced their result.",
    ),
    (
        "Yash Petkar — AS25DXB021",
        "Interface, dashboard and fairness reporting",
        [
            "Clinical console: worklist, image viewer with window and level controls, and "
            "the tabbed analysis surface covering findings, explainability, progression, "
            "uncertainty and the drafted report.",
            "The confidence-as-chroma design system, in which colour saturation encodes "
            "model certainty and drains to an achromatic hatch at the abstention threshold.",
            "Analytics dashboard with cross-filtering, the model-health panel, and the "
            "audit record with sorting, search and export.",
            "The fairness surface: disaggregated equalised-odds reporting, including the "
            "presentation of a breach as a breach.",
            "Accessibility and progressive enhancement: reduced-motion support, keyboard "
            "navigation, a full light theme, and a 3D hero that degrades to 2D without WebGL.",
        ],
        "Judgement call: chose a ranked bar list over a radar chart for the pathology "
        "profile, because radar implies its axes are commensurable and ordered, and "
        "fourteen unrelated pathologies are neither — accepting a less striking chart in "
        "exchange for one that does not imply a false relationship.",
    ),
]

for i, (name, role, items, judgement) in enumerate(MEMBERS):
    para(name, size=13, bold=True, colour=INSTRUMENT, after=1)
    para(role, size=10, italic=True, colour=MUTED, after=6)
    bullets(items)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    r = p.add_run(judgement)
    r.font.size = Pt(9.5)
    r.italic = True
    r.font.color.rgb = MUTED
    if i < len(MEMBERS) - 1:
        doc.add_paragraph()

doc.add_paragraph()
para("Shared work", size=13, bold=True, colour=INSTRUMENT, after=4)
bullets([
    "The design specification, evaluation protocol and abstention thresholds were agreed "
    "jointly before implementation, so that no result could be selected after the fact.",
    "The ethical analysis, including the decision to report the fairness breach rather "
    "than adjust the tolerance to accommodate it.",
    "The report, both presentations and the viva preparation.",
])

doc.add_paragraph()
para("Declaration", size=13, bold=True, colour=INSTRUMENT, after=4)
para(
    "We confirm that this submission is our own work. The deployed classifier uses "
    "publicly released TorchXRayVision weights, which is stated in the report, the model "
    "card, the repository and the application itself; no result is presented as the "
    "outcome of our own training run. All measured figures were produced by running the "
    "system on public NIH ChestX-ray14 data and are reported as measured, including the "
    "two that do not meet the targets we set ourselves.",
    size=10,
    after=18,
)

t = doc.add_table(rows=2, cols=3)
t.style = "Table Grid"
for i, (nm, sid) in enumerate([
    ("Krishna Mathur", "AS25DXB018"),
    ("Atharva Soundankar", "AS25DXB020"),
    ("Yash Petkar", "AS25DXB021"),
]):
    c = t.rows[0].cells[i]
    c.text = ""
    r = c.paragraphs[0].add_run(f"{nm}\n{sid}")
    r.font.size = Pt(9)
    t.rows[1].cells[i].text = ""
    r2 = t.rows[1].cells[i].paragraphs[0].add_run("\nSignature / date\n")
    r2.font.size = Pt(8)
    r2.font.color.rgb = MUTED

doc.add_paragraph()
para(
    "Research prototype. Not a medical device. Must not be used for clinical decisions.",
    size=8.5, italic=True, colour=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER,
)

path = OUT / "SENTINEL-CXR_Contribution_Statement.docx"
doc.save(path)
print(f"Wrote {path.name}")
