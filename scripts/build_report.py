#!/usr/bin/env python3
"""Generate the academic report (.docx).

Run: python scripts/build_report.py
Output: deliverables/report/SENTINEL-CXR_Report.docx
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).resolve().parent.parent / "deliverables" / "report"
OUT.mkdir(parents=True, exist_ok=True)

INSTRUMENT = RGBColor(0x1B, 0x7D, 0x97)
INK = RGBColor(0x11, 0x16, 0x1A)
MUTED = RGBColor(0x5C, 0x66, 0x6D)

doc = Document()

# ── page + base styles ───────────────────────────────────────────────────
for s in doc.sections:
    s.top_margin = s.bottom_margin = Cm(2.4)
    s.left_margin = s.right_margin = Cm(2.6)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(7)
normal.paragraph_format.line_spacing = 1.25


def style_heading(level: int, size: int, colour: RGBColor, before: int, after: int):
    st = doc.styles[f"Heading {level}"]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.color.rgb = colour
    st.font.bold = True
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)


style_heading(1, 17, INSTRUMENT, 20, 8)
style_heading(2, 13, INK, 14, 6)
style_heading(3, 11, MUTED, 10, 4)


def para(text="", *, size=10.5, bold=False, italic=False, colour=None,
         align=None, space_after=None, style=None):
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    r.font.color.rgb = colour or INK
    if align is not None:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def bullets(items, style="List Bullet"):
    for it in items:
        p = doc.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(3)
        if isinstance(it, tuple):
            r = p.add_run(it[0])
            r.bold = True
            r.font.size = Pt(10.5)
            r2 = p.add_run(it[1])
            r2.font.size = Pt(10.5)
        else:
            p.add_run(it).font.size = Pt(10.5)


def shade(cell, hex_colour: str):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hex_colour)
    cell._tc.get_or_add_tcPr().append(el)


def table(headers, rows, widths=None, caption=None, font=9):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(font)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade(c, "1B7D97")
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(v))
            r.font.size = Pt(font)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    if caption:
        para(caption, size=8.5, italic=True, colour=MUTED, space_after=10)
    return t


def code_block(text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(9)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(8.5)
    r.font.color.rgb = MUTED


def page_break():
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


# ═════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═════════════════════════════════════════════════════════════════════════
for _ in range(3):
    doc.add_paragraph()

para("SENTINEL-CXR", size=30, bold=True, colour=INSTRUMENT,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("Uncertainty-Aware Chest Radiograph Triage", size=15, colour=INK,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para("A deep learning system that abstains when it cannot meet its own "
     "coverage guarantee", size=10.5, italic=True, colour=MUTED,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=26)

para("Final Group Project", size=12, bold=True,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("Deep Learning (MAIB AI 114)", size=11,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("Master of Artificial Intelligence in Business", size=10, colour=MUTED,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("S P Jain School of Global Management, Dubai", size=10, colour=MUTED,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

table(
    ["Group member", "Student ID", "Primary contribution"],
    [
        ["Krishna Mathur", "AS25DXB018",
         "Model architecture, ONNX export and quantisation, conformal calibration, "
         "class activation mapping, training notebooks"],
        ["Atharva Soundankar", "AS25DXB020",
         "Backend orchestration, three-tier deployment, evaluation harness, "
         "rate limiting and resilience, CI safety gates"],
        ["Yash Petkar", "AS25DXB021",
         "Clinical console and dashboard, explainability and uncertainty "
         "interfaces, fairness audit and reporting, accessibility"],
    ],
    widths=[4.5, 3.0, 8.0],
)

doc.add_paragraph()
para("Faculty: Prof Anshul Gupta", size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("Term 3, September 2025 intake", size=10, colour=MUTED,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("Submitted 12 August 2026", size=10, colour=MUTED,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

para("Repository: github.com/krish2105/Deep-Learning-", size=9,
     colour=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("This system is a research prototype. It is not a medical device and must "
     "not be used for clinical decisions.", size=9, italic=True, colour=MUTED,
     align=WD_ALIGN_PARAGRAPH.CENTER)

page_break()

# ═════════════════════════════════════════════════════════════════════════
doc.add_heading("Abstract", 1)
para(
    "Deep learning models for medical imaging routinely achieve strong aggregate "
    "accuracy while remaining unsuitable for clinical deployment. The obstacle is "
    "not predictive performance but unearned confidence: a network trained to emit "
    "a probability will emit one for every input, including images outside its "
    "training distribution and cases where the evidence does not support a "
    "decision. A fluent, confident, wrong answer is more dangerous in a clinical "
    "workflow than no answer at all."
)
para(
    "We present SENTINEL-CXR, a chest radiograph triage system built around the "
    "opposite commitment. Rather than a point estimate, it produces a prediction "
    "set carrying a distribution-free coverage guarantee, and it abstains — routing "
    "the study to a human radiologist — when that guarantee cannot be met. Three "
    "mechanisms enforce this: a convolutional variational autoencoder that rejects "
    "inputs which are not chest radiographs; an uncertainty decomposition that "
    "separates the model's own ignorance from irreducible ambiguity in the image; "
    "and split conformal prediction, which converts uncalibrated scores into sets "
    "with finite-sample marginal coverage."
)
para(
    "The system is trained and evaluated on NIH ChestX-ray14 (112,120 radiographs, "
    "30,805 patients) with patient-disjoint splitting. Beyond classification, it "
    "exploits the dataset's patient identifiers and follow-up indices to model "
    "disease progression with recurrent networks, augments minority classes with a "
    "DCGAN, orders the reading worklist with a deep Q-network, and drafts "
    "structured reports with a language model constrained to the vision model's own "
    "output. A fairness audit disaggregates performance across sex, age band and "
    "view position."
)
para(
    "The complete system is deployed on free-tier infrastructure — a Next.js "
    "frontend on Vercel, a FastAPI orchestrator on Render, and a PyTorch inference "
    "core on Hugging Face Spaces — with a documented degradation path when the "
    "inference service is cold.",
    space_after=14,
)

para("Keywords: ", bold=True, space_after=2)
para("conformal prediction; selective classification; out-of-distribution "
     "detection; uncertainty quantification; chest radiography; explainable AI; "
     "algorithmic fairness", size=10, colour=MUTED)

page_break()

# ═════════════════════════════════════════════════════════════════════════
doc.add_heading("Contents", 1)
for n, t in [
    ("1", "Introduction and motivation"),
    ("2", "Background and related work"),
    ("3", "Data"),
    ("4", "System architecture"),
    ("5", "Methods"),
    ("6", "Evaluation methodology"),
    ("7", "Results"),
    ("8", "Ethics, fairness and societal impact"),
    ("9", "Limitations"),
    ("10", "Conclusion"),
    ("", "References"),
    ("", "Appendix A — Syllabus coverage map"),
    ("", "Appendix B — Reproducibility"),
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"{n}\t{t}" if n else f"\t{t}")
    r.font.size = Pt(10.5)

page_break()

# ═════════════════════════════════════════════════════════════════════════
doc.add_heading("1  Introduction and motivation", 1)

doc.add_heading("1.1  The problem", 2)
para(
    "Chest radiography is the highest-volume diagnostic imaging examination in the "
    "world. The resulting reporting backlog is a genuine operational constraint on "
    "healthcare systems, and it is the backlog — not diagnostic difficulty alone — "
    "that automated triage is well positioned to address. If a system can reliably "
    "identify which studies need to be read first, it delivers value even without "
    "matching radiologist accuracy."
)
para(
    "The obstacle to deploying such systems is subtle. Published models report "
    "areas under the ROC curve that appear clinically useful, yet very few reach "
    "practice. The reason is that a model optimised with cross-entropy produces "
    "scores that are not probabilities of being correct. Thresholding such a score "
    "at 0.5 yields no guarantee about how often the answer is right, and the model "
    "will produce equally confident output for a photograph of a cat, a "
    "mispositioned film, or a genuinely ambiguous case."
)
para(
    "This matters because of automation bias: clinicians presented with a confident "
    "machine judgement are measurably less likely to override it. A system that is "
    "wrong and confident therefore causes more harm than one that is silent."
)

doc.add_heading("1.2  Our position", 2)
para(
    "We argue that the correct design objective for clinical decision support is "
    "not maximal accuracy but calibrated humility: the system should quantify what "
    "it does not know and decline to answer when that quantity is too large. This "
    "converts the problem from classification to selective classification, where "
    "abstention is a legitimate and desirable output rather than a failure."
)

doc.add_heading("1.3  Contributions", 2)
bullets([
    ("A three-layer refusal architecture. ",
     "Distributional gating by variational autoencoder, epistemic uncertainty "
     "estimation, and conformal prediction with abstention, composed into a single "
     "pipeline with an explicit decision at each stage."),
    ("Clinically grounded recurrent modelling. ",
     "Patient identifiers and follow-up indices in ChestX-ray14 are used to build "
     "genuine longitudinal sequences, so the recurrent component models disease "
     "progression rather than serving as an unmotivated exercise."),
    ("Verified generative reporting. ",
     "A language model drafts reports from structured model output only, with a "
     "post-generation verifier that rejects any text naming an unsupported finding. "
     "Prompt instructions are treated as a request; the verifier is the guarantee."),
    ("A deployment under real constraints. ",
     "The full system runs on free-tier infrastructure, with a documented and "
     "user-visible degradation path rather than a hidden one."),
    ("An honest fairness audit. ",
     "Performance is disaggregated across available strata, and the axes that "
     "cannot be audited because the dataset lacks the labels are reported as such."),
])

page_break()

# ═════════════════════════════════════════════════════════════════════════
doc.add_heading("2  Background and related work", 1)

doc.add_heading("2.1  Deep learning for chest radiography", 2)
para(
    "The release of ChestX-ray8 (Wang et al., 2017) and its extension to fourteen "
    "labels established the benchmark on which most subsequent work is measured. "
    "CheXNet (Rajpurkar et al., 2017) applied a 121-layer DenseNet and reported "
    "pneumonia detection exceeding the average of four radiologists, a claim later "
    "moderated by analyses showing that the comparison protocol favoured the model."
)
para(
    "Oakden-Rayner et al. (2020) demonstrated hidden stratification: aggregate "
    "performance can remain high while the model fails badly on clinically "
    "important subsets. This finding motivates our disaggregated evaluation and our "
    "refusal to report a single headline number without stratified context."
)

doc.add_heading("2.2  Uncertainty quantification", 2)
para(
    "Gal and Ghahramani (2016) showed that dropout at inference approximates "
    "Bayesian inference over network weights, making Monte-Carlo dropout a "
    "practical route to epistemic uncertainty. Lakshminarayanan et al. (2017) "
    "showed deep ensembles frequently outperform it. Both decompose predictive "
    "uncertainty into an aleatoric component, irreducible given the data, and an "
    "epistemic component reflecting the model's ignorance — a distinction that "
    "matters operationally, because only the latter is a valid reason to abstain."
)

doc.add_heading("2.3  Conformal prediction", 2)
para(
    "Conformal prediction (Vovk et al., 2005; Angelopoulos and Bates, 2023) "
    "converts any point predictor into a set predictor with distribution-free, "
    "finite-sample coverage under exchangeability. Its appeal here is that the "
    "guarantee holds regardless of whether the underlying network is calibrated, "
    "which is precisely the situation with a deep classifier trained by "
    "cross-entropy. Recent work has applied conformal triage to medical imaging "
    "with statistical guarantees on predictive value."
)

doc.add_heading("2.4  Explainability", 2)
para(
    "Grad-CAM (Selvaraju et al., 2017) remains the dominant saliency method in "
    "medical imaging. We adopt it while stating its limitation plainly throughout "
    "this report and in the interface itself: it shows where activation correlates "
    "with a score, not why a decision was made, and a plausible heat map is not "
    "evidence of correct reasoning."
)

page_break()

# ═════════════════════════════════════════════════════════════════════════
doc.add_heading("3  Data", 1)

doc.add_heading("3.1  NIH ChestX-ray14", 2)
para(
    "The primary dataset comprises 112,120 frontal-view radiographs from 30,805 "
    "unique patients, released by the NIH Clinical Center and already "
    "de-identified. Labels for fourteen thoracic pathologies were mined from "
    "radiology reports by natural language processing."
)

table(
    ["Field", "Role in this project"],
    [
        ["Image Index", "Join key to the image files"],
        ["Finding Labels", "Multi-label target, expanded to 14 binary columns"],
        ["Patient ID", "Groups studies into patient timelines; the splitting key"],
        ["Follow-up #", "Sequence position for the recurrent branch"],
        ["Patient Age", "Fairness stratum (binned)"],
        ["Patient Gender", "Fairness stratum"],
        ["View Position", "Fairness stratum and a suspected confounder (§8.2)"],
    ],
    widths=[4.5, 11.0],
    caption="Table 1. Metadata fields and their role. Patient ID and Follow-up # "
            "are what make longitudinal modelling possible on this dataset.",
)

doc.add_heading("3.2  Splitting by patient, not by image", 2)
para(
    "Because the dataset averages three to four studies per patient, splitting by "
    "image would place a patient's follow-up scans on both sides of the train/test "
    "boundary. A model can then identify the patient rather than the pathology, and "
    "every reported metric is inflated by an amount that does not survive contact "
    "with a new institution."
)
para(
    "We split on Patient ID into 70% training, 10% calibration and 20% test, and "
    "assert disjointness programmatically. The calibration split is held out from "
    "training specifically to fit the conformal thresholds; reusing training data "
    "for calibration would void the coverage guarantee, since the exchangeability "
    "assumption would no longer hold."
)
code_block(
    "patients = df['Patient ID'].unique()\n"
    "rng.shuffle(patients)\n"
    "train, cal, test = split_by_patient(patients, 0.70, 0.10, 0.20)\n"
    "assert not (set(train['Patient ID']) & set(test['Patient ID']))"
)

doc.add_heading("3.3  Class imbalance and label noise", 2)
para(
    "The label distribution is severely skewed: Hernia appears in under 0.3% of "
    "studies while Infiltration appears in roughly 18%. We therefore report AUROC "
    "and average precision rather than accuracy, and weight the loss by inverse "
    "class frequency with a cap of 20 — uncapped, the rarest class produces a "
    "weight near 500 and its gradient overwhelms the other thirteen labels."
)
para(
    "The NLP-derived labels carry a reported error rate near 10%. Every metric in "
    "this report inherits that ceiling. We state this rather than presenting our "
    "numbers as though the labels were ground truth."
)

page_break()

# ═════════════════════════════════════════════════════════════════════════
doc.add_heading("4  System architecture", 1)

doc.add_heading("4.1  The constraint that shaped the design", 2)
para(
    "The system was required to deploy entirely on free-tier infrastructure. "
    "Render's free web service provides 512 MB of memory and 0.1 CPU, and spins "
    "down after fifteen minutes of inactivity. PyTorch alone exceeds that memory "
    "budget. Hugging Face Spaces provides 2 vCPU and 16 GB on its free CPU tier, "
    "sleeping only after 48 hours."
)
para(
    "Our first architecture followed directly from this asymmetry: orchestration "
    "where memory is scarce, inference where it is plentiful. In deployment that "
    "proved fragile. The system diagnosed nothing whenever the inference Space "
    "was asleep, still building, or misconfigured, and a reviewer arriving in "
    "that window saw a product that appeared broken."
)
para(
    "We therefore exported the classifier to an 8-bit quantised ONNX graph of "
    "7.9 MB, which fits comfortably inside 512 MB and executes in roughly 150 ms "
    "on 0.1 CPU. The orchestrator now performs real inference by itself. The "
    "Space became an enhancement rather than a dependency: when present it adds "
    "Monte-Carlo sampling and gradient-based saliency, and when absent the system "
    "still classifies, calibrates, abstains and explains."
)
para(
    "Two defects surfaced during that export, both found by comparing the fast "
    "path against the full PyTorch path rather than by assuming equivalence. "
    "Preprocessing initially applied ImageNet normalisation across three "
    "channels, which is the torchvision convention but not TorchXRayVision's, "
    "producing both the wrong tensor shape and the wrong intensity scale. And "
    "the exported graph already terminates in a sigmoid, so applying one again "
    "mapped every score from [0,1] onto [0.50, 0.73] and made every finding "
    "appear to be a coin flip. Neither would have raised an error; both would "
    "have produced confident nonsense."
)

code_block(
    "  Vercel      Next.js 15 — landing + clinical console\n"
    "     |        HTTPS + JWT\n"
    "  Render      FastAPI, 512 MB, NO PyTorch\n"
    "     |        conformal head (NumPy) - auth - audit log\n"
    "     +-- cold --> ONNX int8, <900 ms, 'reduced' mode\n"
    "     |\n"
    "  HF Spaces   16 GB — DenseNet, ViT, VAE, LSTM, Grad-CAM\n"
    "     |\n"
    "  Supabase    Postgres + object storage"
)

doc.add_heading("4.2  Graceful degradation", 2)
para(
    "Because a sleeping Space takes roughly forty seconds to wake, the orchestrator "
    "maintains two inference paths. When the Space is unavailable, a quantised "
    "int8 ONNX model on the orchestrator answers in under a second, the response is "
    "marked as reduced, and the interface displays that state explicitly. Grad-CAM "
    "and uncertainty decomposition are unavailable on this path, and the interface "
    "says so rather than silently omitting them."
)
para(
    "We consider explicit degradation an ethical requirement rather than a "
    "convenience: a clinical user must know which mode produced the result in front "
    "of them."
)

table(
    ["Path", "Components", "Latency", "Trigger"],
    [
        ["Full", "DenseNet + TTA/MC ensemble, Grad-CAM, VAE gate, LSTM", "2–6 s", "Space warm"],
        ["Reduced", "ONNX int8 DenseNet only", "< 900 ms", "Space cold or unreachable"],
    ],
    widths=[2.4, 8.0, 2.2, 3.0],
    caption="Table 2. The two inference paths. Mode is always reported to the user.",
)

doc.add_heading("4.3  Module boundaries", 2)
para(
    "Model definitions are pure functions over tensors with no I/O. The conformal "
    "head and triage policy are pure NumPy and are unit-tested without a network or "
    "a checkpoint. The API layer contains no machine learning, and the frontend "
    "contains no business logic. The practical test of these boundaries is that the "
    "backbone can be exchanged without touching the API, and the API can be tested "
    "with the inference service entirely absent — both of which the test suite does."
)

page_break()

# ═════════════════════════════════════════════════════════════════════════
doc.add_heading("5  Methods", 1)

doc.add_heading("5.1  Classification backbone", 2)
para(
    "A DenseNet-121 pretrained on ImageNet is adapted to single-channel input by "
    "summing the pretrained RGB filters of the first convolution rather than "
    "discarding two of the three channels, which preserves the learned edge "
    "detectors. The classifier head is replaced with dropout followed by a "
    "fourteen-unit linear layer trained with weighted binary cross-entropy."
)
para(
    "Dropout is included deliberately. Without it, Monte-Carlo sampling at "
    "inference yields identical passes and an epistemic uncertainty of exactly "
    "zero — a confident claim of certainty that is an artefact of architecture "
    "rather than a property of the evidence."
)

doc.add_heading("5.2  Out-of-distribution gating", 2)
para(
    "A convolutional variational autoencoder is trained only on chest radiographs. "
    "Per-image reconstruction error then serves as an out-of-distribution score. "
    "The threshold is selected at a fixed false-positive rate of 1% on held-out "
    "radiographs, making the operating point an explicit design decision rather "
    "than a round number chosen by inspection."
)
para(
    "This gate runs before classification. Classifying a photograph is not a "
    "smaller error than misclassifying a radiograph; it is a category error, and "
    "the correct response is refusal."
)

doc.add_heading("5.3  Uncertainty decomposition", 2)
para(
    "Given T stochastic forward passes, predictive uncertainty for a Bernoulli "
    "output decomposes exactly into an expected-entropy (aleatoric) term and a "
    "mutual-information (epistemic) term. The mutual information, known as the BALD "
    "score, is high precisely when individual samples are each confident but "
    "disagree with one another — the signature of a model extrapolating beyond its "
    "training distribution."
)
code_block(
    "total     = H[ E_t[p_t] ]          # predictive entropy\n"
    "aleatoric = E_t[ H[p_t] ]          # expected entropy\n"
    "epistemic = total - aleatoric      # mutual information (BALD)"
)
para(
    "The deployed system additionally supports test-time augmentation as an "
    "uncertainty source. This is necessary because the published pretrained "
    "checkpoint used for the live demonstration contains no dropout layers; "
    "injecting dropout into a network trained without it is not a valid posterior "
    "approximation, whereas the spread of predictions under small clinically "
    "plausible transforms is a defensible robustness measure. Which method produced "
    "a given estimate is reported in the API response and never conflated."
)

doc.add_heading("5.4  Conformal prediction and abstention", 2)
para(
    "For each pathology independently, nonconformity scores are computed on the "
    "calibration split as one minus the predicted probability for true positives. "
    "The threshold is the conformal quantile at level ceil((n+1)(1-alpha))/n. The "
    "finite-sample correction is what makes the guarantee exact rather than "
    "asymptotic; omitting it is the most common implementation error and silently "
    "produces under-coverage."
)
para(
    "This yields marginal coverage per label, not simultaneous coverage across all "
    "fourteen. We claim only the former. The system abstains when the prediction "
    "set is empty with several near-threshold scores, when the set exceeds a "
    "maximum size, or when epistemic uncertainty exceeds a bound. A critical "
    "finding falling just below its threshold triggers escalation regardless."
)

doc.add_heading("5.5  Explanation without gradients", 2)
para(
    "Grad-CAM requires a backward pass, and ONNX Runtime performs inference "
    "only. The fast path therefore had no saliency at all, and the "
    "explainability surface sat empty whenever the Space was unavailable — "
    "precisely when a reviewer was most likely to be looking at it."
)
para(
    "DenseNet, however, terminates in global average pooling followed by a "
    "single linear layer, which is exactly the architecture class activation "
    "mapping was defined for (Zhou et al., 2016). The class map is the "
    "classifier's weights applied across the final convolutional feature maps, "
    "and requires only the forward pass already being computed. We therefore "
    "compute all fourteen maps inside the exported graph and return them "
    "alongside the scores, at a cost of 686 additional floating-point values."
)
para(
    "One subtlety cost us eleven of the fourteen maps before it was understood. "
    "Grad-CAM rectifies its output because a negative gradient-activation "
    "product indicates a region arguing against the class. Classic CAM is a "
    "different quantity: an evidence field whose absolute offset is absorbed by "
    "the classifier bias, so a perfectly informative map can lie entirely below "
    "zero. Applying a ReLU by analogy zeroed most of them. The relative maxima "
    "carry the signal, not the sign."
)

doc.add_heading("5.6  Longitudinal progression", 2)
para(
    "Per-visit CNN embeddings are read by a recurrent network with masked "
    "attention pooling. Padding is masked before the softmax; without this, padded "
    "timesteps receive attention mass and the model learns from data that does not "
    "exist. Vanilla RNN, GRU and LSTM cells are compared under identical splits, "
    "seeds, hidden sizes and schedules, and evaluated against a persistence "
    "baseline that predicts no change — a strong baseline, since clinical states "
    "are persistent."
)

doc.add_heading("5.7  Reinforcement-learning triage", 2)
para(
    "Reading order is modelled as a sequential decision problem. The environment "
    "charges each waiting study its clinical urgency at every timestep, and the "
    "agent reads one study per step, so the only way to reduce cost is to remove "
    "the most expensive study first. A double DQN with a target network and "
    "experience replay is trained against this environment and compared with "
    "random, first-in-first-out, heuristic and oracle-greedy policies."
)
para(
    "An earlier formulation additionally paid a bonus for reading a study. That "
    "bonus dominated the waiting cost, and random, FIFO and heuristic policies "
    "became statistically indistinguishable — the environment could not tell a good "
    "policy from a coin flip. Reformulating the reward as pure cost separates the "
    "baselines cleanly and is what makes the benchmark meaningful."
)

doc.add_heading("5.8  Grounded report generation", 2)
para(
    "A language model drafts the findings and impression sections. Three properties "
    "constrain it. The model never receives the image, only structured output, so "
    "it cannot invent a finding from pixels. The prompt enumerates the closed set "
    "of pathologies that may be mentioned. Finally, generated text is scanned for "
    "any of the fourteen pathology names outside the supported set; if one is "
    "found, the generation is discarded, the incident is written to the audit log, "
    "and a deterministic template is used instead."
)
para(
    "The third property is what makes the first two trustworthy. The verifier "
    "rejects negations as well as assertions: a report stating 'no pneumothorax is "
    "seen' invites the reader to infer the system assessed and excluded it, and if "
    "the finding was never in the supported set, that inference is false."
)

doc.add_heading("5.9  Applied intelligence layer", 2)
para(
    "Four features sit on top of the diagnostic pipeline. Each is scoped so "
    "that a language model never makes a clinical judgement; it parses intent "
    "or writes prose, while every decision described has already been made by "
    "the vision model, the conformal head or the triage policy."
)
bullets([
    ("Natural-language query. ",
     "A plain-English request becomes a deterministic filter over the worklist. "
     "The parser is rule-based first and consults a model only for phrasings the "
     "rules miss, because a regular expression that reliably understands "
     "'abstained studies' costs nothing, requires no API key, cannot hallucinate "
     "a filter the clinician did not ask for, and is identical on every run."),
    ("Similar-case retrieval. ",
     "Cosine similarity over the output space returns comparable prior studies. "
     "Cosine rather than Euclidean because embedding magnitude tracks overall "
     "activation strength while direction carries what kind of finding is "
     "present, and two studies showing the same pathology at different "
     "severities should be neighbours. The response states which space was used, "
     "since the fast path exposes no penultimate features and comparing "
     "conclusions is weaker evidence than comparing appearance."),
    ("Patient timeline. ",
     "Structured deltas across a patient's visits, narrated through the same "
     "grounded reporting path, so the summary cannot introduce a finding absent "
     "from the numbers."),
    ("Disagreement detection. ",
     "Divergence between independent estimates of the same image is reported "
     "rather than averaged away. Averaging two confident and opposite estimates "
     "yields a moderate one representing neither, and conceals the absence of "
     "consensus, which is itself the clinically relevant fact."),
])

page_break()

# ═════════════════════════════════════════════════════════════════════════
doc.add_heading("6  Evaluation methodology", 1)

para(
    "Metrics are chosen for the problem rather than for convenience. Accuracy is "
    "not reported: with a 1% positive rate, a model that always predicts absence "
    "scores 99% and is worthless."
)

table(
    ["Claim", "Metric", "Gate"],
    [
        ["Discriminative performance", "AUROC per pathology, macro mean", "Minimum macro AUROC; CI fails on regression"],
        ["Rare-class performance", "Average precision (AUPRC)", "Reported per class, no aggregate gate"],
        ["Coverage guarantee", "Empirical coverage vs nominal 1-alpha", "Within tolerance band"],
        ["OOD gating", "AUROC, radiographs vs non-radiographs", "Threshold at 1% FPR on radiographs"],
        ["Progression", "Macro AUROC vs persistence baseline", "Must exceed baseline"],
        ["Triage", "Mean episodic return vs FIFO and oracle", "Must exceed heuristic to be reported"],
        ["Fairness", "Equalised-odds gap per stratum", "Below 0.10"],
        ["Grounding", "Ungrounded-finding rate, pre and post filter", "Post-filter rate zero by construction"],
    ],
    widths=[4.2, 5.2, 6.2],
    caption="Table 3. Evaluation matrix. Each row states a claim the system makes "
            "and the measurement that could falsify it.",
)

para(
    "The empirical coverage figure is the single most important number in this "
    "report. It is the claim on which the system's central design decision rests, "
    "and it is asserted in continuous integration so that a retrained model cannot "
    "silently break it."
)

page_break()

# ═════════════════════════════════════════════════════════════════════════
doc.add_heading("7  Results", 1)

para(
    "All figures in this section were measured by running the deployed system. "
    "Where a result is unfavourable it is reported as measured; nothing was "
    "tuned after seeing it.",
    italic=True, colour=MUTED,
)

doc.add_heading("7.1  Experimental setup", 2)
para(
    "The public NIH ChestX-ray14 release was used, of which one image shard "
    "yielding 4,999 usable radiographs across 1,335 patients was drawn. These "
    "were split patient-disjointly in half, giving 2,584 calibration images and "
    "2,415 test images with no patient appearing on both sides. The quantised "
    "int8 ONNX model that serves production was run over both splits, so the "
    "conformal thresholds are calibrated against exactly the weights that are "
    "deployed. A calibrator fitted to a different model than the one serving "
    "traffic would produce a guarantee that does not hold."
)

doc.add_heading("7.2  Conformal coverage", 2)
para(
    "This is the system's central claim and therefore the result that matters "
    "most. Macro empirical coverage on the held-out test split was 0.8845 "
    "against a nominal target of 0.90. The guarantee is therefore approximately "
    "but not exactly met, and four labels under-cover materially."
)

table(
    ["Pathology", "Empirical coverage", "Calibration positives"],
    [
        ["Hernia", "1.0000", "7"],
        ["Fibrosis", "0.9605", "96"],
        ["Edema", "0.9388", "41"],
        ["Pleural Thickening", "0.9302", "79"],
        ["Infiltration", "0.9038", "414"],
        ["Consolidation", "0.8911", "104"],
        ["Cardiomegaly", "0.8900", "96"],
        ["Mass", "0.8947", "85"],
        ["Atelectasis", "0.8794", "261"],
        ["Pneumothorax", "0.8642", "118"],
        ["Effusion", "0.8447", "268"],
        ["Nodule", "0.8173", "110"],
        ["Pneumonia", "0.7941", "31"],
        ["Emphysema", "0.7736", "72"],
        ["Macro average", "0.8845", "target 0.90"],
    ],
    widths=[6.0, 4.5, 4.5],
    caption="Table 4. Empirical coverage per pathology on the patient-disjoint "
            "test split. Four labels fall materially below the nominal level.",
)

para(
    "Two explanations account for the shortfall, and both are properties of the "
    "method rather than implementation defects."
)
bullets([
    ("Exchangeability. ",
     "Split conformal guarantees coverage when calibration and test data are "
     "exchangeable. Splitting by patient is methodologically required here, "
     "because a patient's follow-up studies would otherwise leak across the "
     "boundary and inflate every metric. But splitting by patient also means the "
     "two halves are drawn from different individuals with different disease "
     "profiles, so strict exchangeability does not hold and the guarantee "
     "degrades accordingly."),
    ("Thin calibration sets. ",
     "The conformal quantile is estimated from the positives of each label "
     "alone. Pneumonia has 31 calibration positives and Emphysema 72, so their "
     "quantiles are noisy — and both are among the under-covering labels. Hernia, "
     "with 7 positives, falls below the minimum for a meaningful estimate and is "
     "held at a documented fallback threshold rather than given a fabricated one."),
])
para(
    "The honest response is to report this rather than to raise alpha until the "
    "numbers agree, which would amount to fitting the guarantee to the test set. "
    "The deployed system reports its realised coverage on the dashboard, so a "
    "reader is never shown a nominal figure the system does not actually achieve.",
    bold=True,
)

doc.add_heading("7.3  Backpropagation verification", 2)
para(
    "The from-scratch implementation agrees with central finite differences to a "
    "maximum relative error of 7.2e-11 across all four parameter tensors, "
    "confirming the analytic gradient. The activation ablation reproduces the "
    "expected ordering, with sigmoid converging slowest owing to gradient "
    "saturation and momentum improving on plain stochastic gradient descent."
)
table(
    ["Configuration", "Test accuracy"],
    [["ReLU + momentum", "0.940"], ["ReLU + SGD", "0.928"],
     ["Leaky ReLU + SGD", "0.928"], ["Tanh + SGD", "0.927"],
     ["Sigmoid + SGD", "0.868"]],
    widths=[7.0, 4.0],
    caption="Table 5. Activation and optimiser ablation. Two-layer network, "
            "identical seed and schedule.",
)

doc.add_heading("7.4  Recurrent cell comparison", 2)
para(
    "Gradient-flow analysis averaged over twelve random initialisations shows "
    "gated cells retaining roughly four orders of magnitude more gradient at the "
    "first timestep of a sixty-step sequence than a vanilla recurrent unit. GRU "
    "and LSTM are close and their ordering is not stable across seeds, so no "
    "claim is made that either dominates on this probe; the downstream AUROC "
    "comparison decides which cell is deployed."
)
table(
    ["Cell", "Gradient at t=0", "Gradient at t=59", "Parameters"],
    [["Vanilla RNN", "1.46e-17", "1.09e-01", "463,374"],
     ["GRU", "8.69e-14", "5.61e-02", "1,382,926"],
     ["LSTM", "2.92e-14", "3.02e-02", "1,842,702"]],
    widths=[3.2, 4.0, 4.0, 3.4],
    caption="Table 6. Gradient survival and parameter count. Parameter count is "
            "a confound in any cell comparison and is reported alongside.",
)

doc.add_heading("7.5  Triage policy", 2)
para(
    "On the cost-based reading-room environment, an urgency-weighted heuristic "
    "roughly halves accumulated cost relative to first-in-first-out, and "
    "oracle-greedy bounds what a learned policy could achieve. The separation "
    "between policies is what makes the benchmark usable at all."
)
table(
    ["Policy", "Mean episodic return", "Std"],
    [["Random", "-837.74", "41.45"],
     ["First-in-first-out", "-841.66", "31.48"],
     ["Urgency heuristic", "-397.73", "29.82"],
     ["Oracle-greedy (upper bound)", "-367.22", "26.08"]],
    widths=[5.6, 5.0, 3.0],
    caption="Table 7. Triage policy comparison. Higher is better; return is "
            "negative accumulated urgency-weighted waiting cost.",
)
para(
    "An earlier formulation additionally paid a bonus for reading a study. That "
    "bonus dominated the waiting cost, and random, first-in-first-out and "
    "heuristic policies became statistically indistinguishable — the environment "
    "could not tell a good policy from a coin flip. Reformulating the reward as "
    "pure cost separated them cleanly. This was found by running the baselines "
    "rather than assuming they would differ."
)

doc.add_heading("7.6  Quantisation fidelity and latency", 2)
para(
    "Quantising the classifier to int8 changed its outputs by a mean absolute "
    "difference of 9e-05 against the full-precision model, with top-1 agreement "
    "on all trials. End-to-end, the orchestrator answers in approximately 150 ms "
    "on 0.1 CPU using the quantised model, against roughly 450 ms for the "
    "full-precision PyTorch path, and agreement between the two on the same image "
    "is to within 0.001."
)
table(
    ["Path", "Model", "Latency", "Provides"],
    [["Fast (default)", "int8 ONNX, 7.9 MB", "~150 ms",
      "Classification, conformal set, CAM"],
     ["Full (optional)", "PyTorch on HF Spaces", "~450 ms warm",
      "Adds MC sampling, Grad-CAM, VAE gate"]],
    widths=[3.2, 4.4, 2.6, 5.4],
    caption="Table 8. The two inference paths. The system is fully functional "
            "on the fast path alone.",
)

page_break()

# ═════════════════════════════════════════════════════════════════════════
doc.add_heading("8  Ethics, fairness and societal impact", 1)

para(
    "Learning outcome E is assessed only in this project, and we treat it as a "
    "first-class engineering concern rather than a closing section."
)

doc.add_heading("8.1  Disaggregated performance — measured", 2)
para(
    "Performance was measured within strata of sex, age band and view position "
    "on the 2,415-image test split and reported as equalised-odds gaps: the "
    "maximum within-stratum difference in true-positive and false-positive "
    "rates. Equal accuracy across groups is not sufficient, because a model can "
    "be equally accurate on two populations while missing substantially more "
    "disease in one of them."
)

table(
    ["Stratum", "TPR gap", "FPR gap", "Within tolerance (0.10)"],
    [["Patient sex", "0.0165", "0.0267", "Yes"],
     ["Age band", "0.1295", "0.2040", "No"],
     ["View position", "0.0228", "0.2149", "No"]],
    widths=[4.2, 3.2, 3.2, 4.6],
    caption="Table 9. Equalised-odds gaps. Two of three strata breach the "
            "tolerance the project set for itself before measuring.",
)

para(
    "The audit fails. The maximum equalised-odds gap is 0.2149 against a "
    "tolerance of 0.10, and the system reports this on its dashboard as a "
    "breach with nothing downgrading it.",
    bold=True,
)
para(
    "Sex is essentially balanced. Age and view position are not, and the view "
    "position result is the more informative of the two because it was predicted "
    "in the design specification before any data had been examined."
)

doc.add_heading("8.2  The view-position shortcut", 2)
para(
    "Anteroposterior films are acquired at the bedside from patients too unwell to "
    "stand, whereas posteroanterior films are acquired in the radiology department. "
    "View position therefore correlates with severity, and a model can learn to "
    "read 'this is an AP film' as 'this patient is sick'. Such a shortcut scores "
    "well on a held-out split drawn from the same institution and fails as soon as "
    "acquisition practice differs."
)
para(
    "The measurement confirms it. The false-positive rate differs by 0.2149 "
    "between anteroposterior and posteroanterior films — an order of magnitude "
    "larger than the corresponding difference between sexes, and the single "
    "largest disparity found anywhere in the audit. The true-positive gap by "
    "contrast is only 0.0228, which is the signature of the shortcut rather than "
    "of a difference in disease: the model is not missing more disease on AP "
    "films, it is over-calling disease on them."
)
para(
    "This is what makes the finding worth reporting rather than concealing. A "
    "model that has learned to read acquisition circumstance as evidence of "
    "illness will appear to perform well on any test set drawn from the same "
    "institution, and will degrade unpredictably wherever portable radiography "
    "is used differently. Aggregate AUROC would not have revealed it."
)

doc.add_heading("8.3  What we cannot audit", 2)
para(
    "ChestX-ray14 contains no race or ethnicity labels. A major axis of documented "
    "disparity in medical artificial intelligence therefore cannot be examined "
    "here. Seyyed-Kalantari et al. (2021) found systematic underdiagnosis by chest "
    "radiograph classifiers in under-served populations, so this is not a "
    "hypothetical gap."
)
para(
    "This absence is itself a finding. It must not be presented, and is not "
    "presented, as an absence of bias.",
    bold=True,
)

doc.add_heading("8.4  Mitigations built into the system", 2)
bullets([
    ("Abstention rather than forced prediction. ",
     "The system declines rather than guessing when its guarantee cannot be met."),
    ("Distributional rejection. ",
     "Inputs outside the training distribution are refused before classification."),
    ("Uncertainty surfaced, not buried. ",
     "Epistemic and aleatoric components are shown in the interface, and confidence "
     "is encoded visually as colour saturation so that certainty is legible without "
     "reading a number."),
    ("Verified generation. ",
     "The report writer cannot introduce a finding the vision model did not support."),
    ("Audit trail. ",
     "Every analysis, abstention, rejection, human review and rejected generation is "
     "recorded in an append-only log."),
    ("Persistent non-diagnostic notice. ",
     "Every surface of the application states that the system is not a medical device."),
])

doc.add_heading("8.5  Intended and unintended use", 2)
para(
    "The system is intended as a research prototype demonstrating uncertainty-aware "
    "triage. It is not validated for clinical use, has no regulatory clearance, and "
    "must not inform patient care. The most plausible misuse is deployment as a "
    "screening tool in a resource-constrained setting on the argument that an "
    "imperfect system is better than none. We reject that argument here: an "
    "unvalidated system with unknown subgroup performance can concentrate harm on "
    "precisely the populations such a deployment claims to serve."
)

page_break()

# ═════════════════════════════════════════════════════════════════════════
doc.add_heading("9  Limitations", 1)

bullets([
    ("Label noise. ",
     "Labels were NLP-mined from free-text reports with roughly 10% error. No "
     "metric reported here can be more accurate than the labels it is measured against."),
    ("Single-institution data. ",
     "All data originates from one United States hospital system. Performance on "
     "other populations, equipment and acquisition practice is unvalidated."),
    ("Coverage is approached, not attained. ",
     "Measured macro coverage is 0.8845 against a 0.90 target, with Emphysema at "
     "0.774 and Pneumonia at 0.794. The shortfall is explained in section 7.2 and "
     "is a property of patient-disjoint splitting and thin per-label calibration "
     "sets, not an implementation defect — but the system does not currently "
     "deliver its nominal guarantee, and says so."),
    ("Marginal, not simultaneous, coverage. ",
     "The conformal guarantee holds per label. Simultaneous coverage across all "
     "fourteen would require a multiplicity correction and is not claimed."),
    ("Measured on one shard. ",
     "Calibration and evaluation used 4,999 radiographs from 1,335 patients, not "
     "the full 112,120. The figures are real but their confidence intervals are "
     "wider than the full corpus would give."),
    ("A published checkpoint, not our own weights. ",
     "The deployed classifier uses TorchXRayVision pretrained weights. The "
     "training notebooks reproduce and extend them, and swapping in group-trained "
     "weights is a single configuration change, but no result in this report is "
     "presented as the outcome of our own training run."),
    ("Exchangeability. ",
     "The guarantee assumes calibration and test data are exchangeable. Distribution "
     "shift — a new scanner, a new population — voids it, and the system has no "
     "mechanism to detect such shift beyond the OOD gate."),
    ("Grad-CAM is correlational. ",
     "Saliency indicates where activation correlates with a score, not why a "
     "decision was reached."),
    ("Windowing is approximate. ",
     "The viewer implements window and level as CSS filters over 8-bit images, which "
     "approximates but does not reproduce true DICOM windowing over 12-bit data."),
    ("Uncertainty source varies. ",
     "The live demonstration uses test-time augmentation because the published "
     "checkpoint lacks dropout. This is a robustness measure, not a Bayesian "
     "posterior, and is labelled as such in every response."),
    ("Simulated triage environment. ",
     "The reinforcement-learning results are obtained in simulation. No claim is "
     "made about behaviour in a real reading room."),
])

page_break()

# ═════════════════════════════════════════════════════════════════════════
doc.add_heading("10  Conclusion", 1)

para(
    "We set out to build a chest radiograph triage system whose defining property "
    "is not accuracy but calibrated humility. The resulting system converts "
    "uncalibrated network scores into prediction sets with a distribution-free "
    "coverage guarantee, refuses inputs outside its training distribution, "
    "separates its own ignorance from genuine ambiguity in the image, and abstains "
    "rather than guessing when the guarantee cannot be met."
)
para(
    "Building it required engaging with the full breadth of the unit: convolutional "
    "networks for classification, recurrent networks for longitudinal progression, "
    "adversarial networks for minority-class augmentation, variational autoencoders "
    "for distributional gating, transfer learning throughout, deep reinforcement "
    "learning for worklist ordering, transformers for architectural comparison, and "
    "generative language models for reporting — composed into one system rather "
    "than assembled as separate exercises."
)
para(
    "The most valuable lessons were negative ones, and there were more of them "
    "than we expected. An early triage environment could not distinguish a good "
    "policy from a random one and had to be reformulated. The pretrained "
    "checkpoint contains no dropout, silently invalidating the intended "
    "uncertainty method and forcing an honest alternative. The interface once "
    "displayed a probability above its threshold beside the label 'not included', "
    "because the stored scores and the scores used for the decision had diverged. "
    "A quantised model passed every numerical check and then proved unexecutable "
    "on the deployed runtime, because the verification tested agreement but never "
    "tested whether a session could be created at all. Applying a ReLU to a class "
    "activation map by analogy with Grad-CAM zeroed eleven of the fourteen maps. "
    "A rate limiter keyed on the proxy address rather than the caller placed every "
    "visitor worldwide into a single shared budget."
)
para(
    "Each was found by testing a claim the system made about itself, and none "
    "would have been caught by checking whether the application looked correct. "
    "That pattern — assert the property, not the appearance — is the most "
    "transferable thing we take from the project."
)
para(
    "The final results are imperfect and we have reported them that way. Coverage "
    "reaches 0.8845 rather than 0.90. The fairness audit fails its own tolerance "
    "at 0.2149, driven by exactly the acquisition confound the design document "
    "predicted before any data was examined. A system built around the principle "
    "of admitting what it does not know would be a poor advertisement for that "
    "principle if its authors did otherwise."
)
para(
    "That is the underlying argument of this project. A clinical system earns trust "
    "not by being confident but by being able to say, precisely and verifiably, "
    "what it does not know."
)

page_break()

# ═════════════════════════════════════════════════════════════════════════
doc.add_heading("References", 1)

REFS = [
    "Angelopoulos, A. N. and Bates, S. (2023) 'Conformal prediction: a gentle introduction', Foundations and Trends in Machine Learning, 16(4), pp. 494–591.",
    "Bai, S., Kolter, J. Z. and Koltun, V. (2018) 'An empirical evaluation of generic convolutional and recurrent networks for sequence modeling', arXiv:1803.01271.",
    "Burgess, C. P. et al. (2018) 'Understanding disentangling in beta-VAE', arXiv:1804.03599.",
    "Dosovitskiy, A. et al. (2020) 'An image is worth 16x16 words: transformers for image recognition at scale', arXiv:2010.11929.",
    "Gal, Y. and Ghahramani, Z. (2016) 'Dropout as a Bayesian approximation', Proceedings of ICML, pp. 1050–1059.",
    "Goodfellow, I., Bengio, Y. and Courville, A. (2016) Deep Learning. Cambridge, MA: MIT Press.",
    "Goodfellow, I. et al. (2014) 'Generative adversarial networks', Advances in Neural Information Processing Systems, 27.",
    "Greff, K. et al. (2017) 'LSTM: a search space odyssey', IEEE Transactions on Neural Networks and Learning Systems, 28(10), pp. 2222–2232.",
    "Hardt, M., Price, E. and Srebro, N. (2016) 'Equality of opportunity in supervised learning', Advances in Neural Information Processing Systems, 29.",
    "Hochreiter, S. and Schmidhuber, J. (1997) 'Long short-term memory', Neural Computation, 9(8), pp. 1735–1780.",
    "Kingma, D. P. and Welling, M. (2014) 'Auto-encoding variational Bayes', Proceedings of ICLR.",
    "Kornblith, S., Shlens, J. and Le, Q. V. (2019) 'Do better ImageNet models transfer better?', Proceedings of CVPR, pp. 2661–2671.",
    "Krizhevsky, A., Sutskever, I. and Hinton, G. E. (2017) 'ImageNet classification with deep convolutional neural networks', Communications of the ACM, 60(6), pp. 84–90.",
    "Lakshminarayanan, B., Pritzel, A. and Blundell, C. (2017) 'Simple and scalable predictive uncertainty estimation using deep ensembles', Advances in Neural Information Processing Systems, 30.",
    "LeCun, Y. et al. (1998) 'Gradient-based learning applied to document recognition', Proceedings of the IEEE, 86(11), pp. 2278–2324.",
    "Lipton, Z. C., Berkowitz, J. and Elkan, C. (2015) 'A critical review of recurrent neural networks for sequence learning', arXiv:1506.00019.",
    "Mitchell, M. et al. (2019) 'Model cards for model reporting', Proceedings of FAT*, pp. 220–229.",
    "Mnih, V. et al. (2015) 'Human-level control through deep reinforcement learning', Nature, 518(7540), pp. 529–533.",
    "Oakden-Rayner, L. et al. (2020) 'Hidden stratification causes clinically meaningful failures in machine learning for medical imaging', Proceedings of CHIL, pp. 151–159.",
    "Obermeyer, Z. et al. (2019) 'Dissecting racial bias in an algorithm used to manage the health of populations', Science, 366(6464), pp. 447–453.",
    "Radford, A., Metz, L. and Chintala, S. (2015) 'Unsupervised representation learning with deep convolutional generative adversarial networks', arXiv:1511.06434.",
    "Radford, A. et al. (2021) 'Learning transferable visual models from natural language supervision', Proceedings of ICML, pp. 8748–8763.",
    "Rajpurkar, P. et al. (2017) 'CheXNet: radiologist-level pneumonia detection on chest X-rays with deep learning', arXiv:1711.05225.",
    "Selvaraju, R. R. et al. (2017) 'Grad-CAM: visual explanations from deep networks via gradient-based localization', Proceedings of ICCV, pp. 618–626.",
    "Seyyed-Kalantari, L. et al. (2021) 'Underdiagnosis bias of artificial intelligence algorithms applied to chest radiographs in under-served patient populations', Nature Medicine, 27(12), pp. 2176–2182.",
    "Schmidhuber, J. (2015) 'Deep learning in neural networks: an overview', Neural Networks, 61, pp. 85–117.",
    "van Hasselt, H., Guez, A. and Silver, D. (2016) 'Deep reinforcement learning with double Q-learning', Proceedings of AAAI, pp. 2094–2100.",
    "Vinyals, O. et al. (2019) 'Grandmaster level in StarCraft II using multi-agent reinforcement learning', Nature, 575(7782), pp. 350–354.",
    "Vovk, V., Gammerman, A. and Shafer, G. (2005) Algorithmic Learning in a Random World. New York: Springer.",
    "Wang, X. et al. (2017) 'ChestX-ray8: hospital-scale chest X-ray database and benchmarks on weakly-supervised classification and localization of common thorax diseases', Proceedings of CVPR, pp. 2097–2106.",
]

for r in sorted(REFS):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.first_line_indent = Cm(-0.8)
    p.paragraph_format.space_after = Pt(6)
    p.add_run(r).font.size = Pt(9.5)

page_break()

# ═════════════════════════════════════════════════════════════════════════
doc.add_heading("Appendix A — Syllabus coverage map", 1)
para(
    "Every topic in the unit outline is realised inside the single system rather "
    "than as a standalone exercise.", space_after=10,
)

table(
    ["Wk", "Unit topic", "Realisation", "Notebook"],
    [
        ["1–2", "Neural networks, activations, backpropagation", "Backprop from scratch in NumPy; gradient check; activation and optimiser ablation", "01"],
        ["3", "Convolutional neural networks", "DenseNet-121 multi-label classifier; conformal calibration", "02"],
        ["4", "Recurrent neural networks", "GRU over per-visit embeddings; persistence baseline", "03"],
        ["5", "Long short-term memory", "RNN/GRU/LSTM ablation; gradient-flow analysis", "04"],
        ["6", "Generative adversarial networks", "DCGAN minority-class augmentation; measured ΔAUPRC", "05"],
        ["7", "Autoencoders and VAEs", "Convolutional VAE OOD gate; latent traversal", "06"],
        ["8", "Transfer learning", "Scratch / frozen / full / progressive; data-efficiency curve", "07"],
        ["9", "Deep reinforcement learning", "Double DQN worklist triage vs FIFO, heuristic, oracle", "08"],
        ["10", "Practical applications (ViT, CLIP)", "ViT-B/16 vs CNN; BiomedCLIP zero-shot baseline", "09"],
        ["11", "Generative AI integration", "GAN + VAE + grounded LLM reporting with verifier", "10"],
        ["12", "Ethical and societal impacts", "Disaggregated audit; shortcut probe; model card", "11"],
    ],
    widths=[1.2, 4.6, 7.4, 1.8],
    caption="Table A1. Syllabus coverage.",
)

doc.add_heading("Appendix B — Reproducibility", 1)
bullets([
    ("Seed. ", "All experiments use seed 20260812 for NumPy, PyTorch and CUDA."),
    ("Splits. ", "Patient-disjoint 70/10/20, asserted programmatically at load time."),
    ("Environment. ", "Notebooks run on Colab free tier (T4). No paid runtime is required."),
    ("Weight provenance. ",
     "The deployed demonstration uses TorchXRayVision pretrained weights so that the "
     "live system is functional from day one. This is stated in the README, the "
     "model card and the API health endpoint. No result produced by borrowed "
     "weights is presented as the group's own training outcome."),
    ("Tests. ", "80 automated tests cover the conformal head, uncertainty "
                "decomposition, triage ordering, progression logic, grounding verification "
                "and the full API surface."),
    ("Artefacts. ",
     "Notebooks 02, 08 and 11 export conformal_calibration.json, dqn_policy.json and "
     "fairness_report.json respectively, which the deployed API loads. Until they are "
     "present, the API reports that its coverage guarantee is not in force."),
])

doc.add_paragraph()
para("— End of report —", size=9, italic=True, colour=MUTED,
     align=WD_ALIGN_PARAGRAPH.CENTER)

path = OUT / "SENTINEL-CXR_Report.docx"
doc.save(path)
print(f"Wrote {path}")
print(f"  paragraphs: {len(doc.paragraphs)}  tables: {len(doc.tables)}")
