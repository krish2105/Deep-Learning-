#!/usr/bin/env python3
"""Render the academic report to PDF.

Reads the generated .docx and re-lays it out with ReportLab, rather than
maintaining a second copy of the text. The .docx generator stays the single
source of truth, so the two documents cannot drift apart — which is the same
reason the notebooks and decks are generated rather than hand-edited.

(LibreOffice would also convert it, but it is not installed here and pulling in
a 400 MB dependency to move text between two formats is not a good trade.)

Run:  python scripts/build_report.py && python scripts/build_report_pdf.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import RGBColor as DocxRGB
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    KeepTogether,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "deliverables" / "report" / "SENTINEL-CXR_Report.docx"
OUT = ROOT / "deliverables" / "report" / "SENTINEL-CXR_Report.pdf"

INSTRUMENT = colors.HexColor("#1B7D97")
INK = colors.HexColor("#11161A")
MUTED = colors.HexColor("#5C666D")
RULE = colors.HexColor("#DFE3E6")

base = getSampleStyleSheet()

S = {
    "title": ParagraphStyle(
        "t", parent=base["Title"], fontName="Helvetica-Bold", fontSize=30,
        textColor=INSTRUMENT, spaceAfter=4, alignment=TA_CENTER, leading=34,
    ),
    "subtitle": ParagraphStyle(
        "st", parent=base["Normal"], fontName="Helvetica", fontSize=14,
        textColor=INK, alignment=TA_CENTER, spaceAfter=6, leading=18,
    ),
    "centre": ParagraphStyle(
        "c", parent=base["Normal"], fontName="Helvetica", fontSize=10,
        textColor=MUTED, alignment=TA_CENTER, spaceAfter=3, leading=14,
    ),
    "h1": ParagraphStyle(
        "h1", parent=base["Heading1"], fontName="Helvetica-Bold", fontSize=16,
        textColor=INSTRUMENT, spaceBefore=16, spaceAfter=7, leading=19,
    ),
    "h2": ParagraphStyle(
        "h2", parent=base["Heading2"], fontName="Helvetica-Bold", fontSize=12,
        textColor=INK, spaceBefore=11, spaceAfter=5, leading=15,
    ),
    "h3": ParagraphStyle(
        "h3", parent=base["Heading3"], fontName="Helvetica-Bold", fontSize=10.5,
        textColor=MUTED, spaceBefore=8, spaceAfter=3, leading=13,
    ),
    "body": ParagraphStyle(
        "b", parent=base["Normal"], fontName="Helvetica", fontSize=9.8,
        textColor=INK, alignment=TA_JUSTIFY, spaceAfter=6, leading=14,
    ),
    "bullet": ParagraphStyle(
        "bu", parent=base["Normal"], fontName="Helvetica", fontSize=9.8,
        textColor=INK, leftIndent=11, bulletIndent=2, spaceAfter=4, leading=13.5,
    ),
    "mono": ParagraphStyle(
        "m", parent=base["Normal"], fontName="Courier", fontSize=8,
        textColor=MUTED, leftIndent=8, spaceAfter=8, leading=10.5,
    ),
    "caption": ParagraphStyle(
        "cap", parent=base["Normal"], fontName="Helvetica-Oblique", fontSize=8,
        textColor=MUTED, spaceAfter=10, leading=11,
    ),
    "ref": ParagraphStyle(
        "r", parent=base["Normal"], fontName="Helvetica", fontSize=8.6,
        textColor=INK, leftIndent=14, firstLineIndent=-14, spaceAfter=5, leading=11.5,
    ),
}


def esc(text: str) -> str:
    return (
        text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    )


def is_mono(p) -> bool:
    return any(r.font.name == "Consolas" for r in p.runs if r.font.name)


def is_caption(p) -> bool:
    return bool(p.runs) and all(r.italic for r in p.runs if r.text.strip())


def build_table(t) -> Table:
    rows = [[esc(c.text) for c in r.cells] for r in t.rows]
    if not rows:
        return None

    n = len(rows[0])
    header_style = ParagraphStyle(
        "th", fontName="Helvetica-Bold", fontSize=7.6, textColor=colors.white, leading=9.5
    )
    cell_style = ParagraphStyle(
        "td", fontName="Helvetica", fontSize=7.6, textColor=INK, leading=9.5
    )

    data = [[Paragraph(c, header_style) for c in rows[0]]] + [
        [Paragraph(c, cell_style) for c in r] for r in rows[1:]
    ]

    avail = 170 * mm
    tbl = Table(data, colWidths=[avail / n] * n, repeatRows=1, hAlign="LEFT")
    tbl.setStyle(
        TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), INSTRUMENT),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("GRID", (0, 0), (-1, -1), 0.4, RULE),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F7F8F9")]),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("LEFTPADDING", (0, 0), (-1, -1), 5),
            ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ])
    )
    return tbl


def on_page(canvas, doc):
    """Footer with page number and the non-diagnostic notice."""
    canvas.saveState()
    canvas.setStrokeColor(RULE)
    canvas.setLineWidth(0.5)
    canvas.line(20 * mm, 15 * mm, 190 * mm, 15 * mm)
    canvas.setFont("Helvetica", 7)
    canvas.setFillColor(MUTED)
    canvas.drawString(
        20 * mm, 10.5 * mm,
        "SENTINEL-CXR · MAIB AI 114 · research prototype, not a medical device",
    )
    canvas.drawRightString(190 * mm, 10.5 * mm, str(canvas.getPageNumber()))
    canvas.restoreState()


def main() -> None:
    if not SRC.exists():
        raise SystemExit(f"{SRC} not found — run scripts/build_report.py first")

    doc = Document(SRC)
    story: list = []

    # Walk body elements in document order so tables stay where they belong.
    from docx.table import Table as DocxTable
    from docx.text.paragraph import Paragraph as DocxPara

    body = doc.element.body
    para_map = {p._p: p for p in doc.paragraphs}
    table_map = {t._tbl: t for t in doc.tables}

    first_heading_seen = False

    for child in body.iterchildren():
        if child in para_map:
            p = para_map[child]
            text = p.text.strip()
            style = p.style.name

            if not text:
                story.append(Spacer(1, 4))
                continue

            if style == "Heading 1":
                # Start each top-level section on a fresh page, except the first.
                if first_heading_seen:
                    story.append(PageBreak())
                first_heading_seen = True
                story.append(Paragraph(esc(text), S["h1"]))
            elif style == "Heading 2":
                story.append(Paragraph(esc(text), S["h2"]))
            elif style == "Heading 3":
                story.append(Paragraph(esc(text), S["h3"]))
            elif style.startswith("List"):
                story.append(Paragraph(esc(text), S["bullet"], bulletText="—"))
            elif is_mono(p):
                for line in text.split("\n"):
                    story.append(Paragraph(esc(line).replace(" ", "&nbsp;"), S["mono"]))
            elif p.paragraph_format.first_line_indent and p.paragraph_format.first_line_indent < 0:
                story.append(Paragraph(esc(text), S["ref"]))
            elif is_caption(p):
                story.append(Paragraph(esc(text), S["caption"]))
            elif not first_heading_seen:
                # Title-page block
                size = p.runs[0].font.size.pt if p.runs and p.runs[0].font.size else 10
                if size >= 28:
                    story.append(Paragraph(esc(text), S["title"]))
                elif size >= 13:
                    story.append(Paragraph(esc(text), S["subtitle"]))
                else:
                    story.append(Paragraph(esc(text), S["centre"]))
            else:
                story.append(Paragraph(esc(text), S["body"]))

        elif child in table_map:
            tbl = build_table(table_map[child])
            if tbl is not None:
                story.append(Spacer(1, 3))
                story.append(tbl)
                story.append(Spacer(1, 7))

    template = BaseDocTemplate(
        str(OUT),
        pagesize=A4,
        title="SENTINEL-CXR — Uncertainty-Aware Chest Radiograph Triage",
        author="Krishna Mathur, Atharva Soundankar, Yash Petkar",
        subject="Final Group Project, Deep Learning (MAIB AI 114)",
        leftMargin=20 * mm, rightMargin=20 * mm,
        topMargin=18 * mm, bottomMargin=20 * mm,
    )
    frame = Frame(
        template.leftMargin, template.bottomMargin,
        template.width, template.height, id="body",
    )
    template.addPageTemplates([PageTemplate(id="main", frames=frame, onPage=on_page)])
    template.build(story)

    size_kb = OUT.stat().st_size / 1024
    print(f"Wrote {OUT}")
    print(f"  {size_kb:.0f} KB")


if __name__ == "__main__":
    main()
